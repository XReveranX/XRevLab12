import PySimpleGUI as sg
from docx import Document
from pokets.Iron import Ironf
from pokets.TV import TVf
from pokets.WM import WNf


def show_enter():
    ironr = 0
    tvr = 0
    wnr = 0

    # All the stuff inside your window.
    layout = [  [sg.Text("Введите время использования утюга (ч)")],
                [sg.InputText()],
                [sg.Text("Введите время использования телевизора (ч)")],
                [sg.InputText()],
                [sg.Text("Введите время использования стиральной машины (ч)")],
                [sg.InputText()],
                [sg.Text("Введите ваш тариф (Руб/кВт*ч)")],
                [sg.InputText()],
                [sg.Button('Ok'), sg.Button('Cancel')] ]

    # Create the Window
    window = sg.Window('programm', layout)

    # Event Loop to process "events" and get the "values" of the inputs
    while True:
        event, values = window.read()

        # if user closes window or clicks cancel
        if event == sg.WIN_CLOSED or event == 'Cancel':
            break

        ironr = Ironf(int(values[0]),(int(values[3])))
        tvr = TVf(int(values[1]),(int(values[3])))
        wnr = WNf(int(values[2]),(int(values[3])))
        break

    window.close()
    return(ironr, tvr, wnr)

def show_result(ironr2, tvr2, wnr2):
    doc = Document()
    ironr2 = str(("Затраты на использование утюга", ironr2, "руб"))
    tvr2 = str(("Затраты на использование телевизора", tvr2, "руб"))
    wnr2 = str(("Затраты на использование стиральной машины", wnr2, "руб"))
    # All the stuff inside your window.
    layout2 = [ [sg.Text(ironr2)],
                [sg.Text(tvr2)],
                [sg.Text(wnr2)],
                [sg.Button('docx')],
                [sg.Button('Cancel')] ]

    # Create the Window
    window2 = sg.Window('programm', layout2)

    # Event Loop to process "events" and get the "values" of the inputs
    while True:
        event2, values2 = window2.read()

        # if user closes window or clicks cancel
        if event2 == sg.WIN_CLOSED or event2 == 'Cancel':
            break
        if event2 == 'docx':
            doc.add_paragraph(ironr2)
            doc.add_paragraph(tvr2)
            doc.add_paragraph(wnr2)
            doc.save("otchet.docx")

    window2.close()

res=[]
res = show_enter()
show_result(res[0], res[1], res[2])
